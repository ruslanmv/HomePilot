"""
Vector database module for HomePilot using ChromaDB
Provides RAG (Retrieval Augmented Generation) capabilities for project knowledge bases
"""
import os
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional

try:
    import chromadb
    from chromadb.config import Settings
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False
    chromadb = None
    Settings = None
    print("Warning: chromadb package not installed. Install with: pip install chromadb")

# Optional document format support (additive)
try:
    import docx as _docx_lib  # python-docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl as _openpyxl_lib
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

from .config import UPLOAD_DIR

# Export for use in other modules
__all__ = ['CHROMADB_AVAILABLE', 'get_chroma_client', 'query_project_knowledge', 'get_project_document_count', 'process_and_add_file', 'collection_name', 'DEFAULT_NAMESPACE']

# ChromaDB persistent storage location
CHROMA_DB_PATH = Path(UPLOAD_DIR) / "chroma_db"

# Initialize ChromaDB client
def get_chroma_client():
    """Get or create ChromaDB client with persistent storage"""
    if not CHROMADB_AVAILABLE:
        raise ImportError("ChromaDB is not installed. Install with: pip install chromadb")

    CHROMA_DB_PATH.mkdir(parents=True, exist_ok=True)

    client = chromadb.PersistentClient(
        path=str(CHROMA_DB_PATH),
        settings=Settings(
            anonymized_telemetry=False,
            allow_reset=True
        )
    )
    return client

#: The namespace every existing caller uses. Named rather than inlined so that the one place
#: the collection name is built is also the one place this default is stated.
DEFAULT_NAMESPACE = "project"


def collection_name(key: str, namespace: str = DEFAULT_NAMESPACE) -> str:
    """The Chroma collection for one key inside one namespace.

    ``namespace="project"`` produces exactly the name this module has always produced —
    ``project_<md5(project_id)[:16]>`` — so every existing collection on every existing install
    is found unchanged. A test asserts that against a fixed hash, because "unchanged" is a
    claim about data already on disk and a wrong one orphans somebody's knowledge base.

    Other namespaces exist so that things which are *not* project knowledge can be stored in
    the same Chroma without being seen by the functions that count, query and delete a
    project's documents. MeetingSense (MS15) uses ``"meetings"``: a meeting is retrieved from,
    never absorbed into a project (D4), and ``get_project_document_count`` reporting a number
    that jumps because somebody recorded a call would be a bug with no visible cause.

    One function rather than an f-string at each site: two copies of a naming rule is exactly
    how a delete stops matching the create it is meant to undo.
    """
    return f"{namespace}_{hashlib.md5(key.encode()).hexdigest()[:16]}"


def get_or_create_collection(project_id: str, namespace: str = DEFAULT_NAMESPACE):
    """Get or create a collection for a specific project (or another namespace)."""
    client = get_chroma_client()

    # Collection name must be alphanumeric + underscores
    name = collection_name(project_id, namespace)

    # The project metadata is left exactly as it was: it is written into collections that
    # already exist on disk, and a differently-shaped dict here would be a second schema for
    # the same records.
    metadata = (
        {"project_id": project_id}
        if namespace == DEFAULT_NAMESPACE
        else {"namespace": namespace, "key": project_id}
    )

    collection = client.get_or_create_collection(
        name=name,
        metadata=metadata
    )

    return collection

def add_documents_to_project(
    project_id: str,
    documents: List[str],
    metadatas: Optional[List[Dict[str, Any]]] = None,
    ids: Optional[List[str]] = None
) -> int:
    """
    Add documents to a project's knowledge base

    Args:
        project_id: Project identifier
        documents: List of document texts
        metadatas: Optional metadata for each document
        ids: Optional IDs for each document (auto-generated if not provided)

    Returns:
        Number of documents added
    """
    if not documents:
        return 0

    collection = get_or_create_collection(project_id)

    # Generate IDs if not provided
    if ids is None:
        ids = [hashlib.md5(doc.encode()).hexdigest() for doc in documents]

    # Generate metadata if not provided
    if metadatas is None:
        metadatas = [{"source": "uploaded_file"} for _ in documents]

    collection.add(
        documents=documents,
        metadatas=metadatas,
        ids=ids
    )

    return len(documents)

def query_project_knowledge(
    project_id: str,
    query: str,
    n_results: int = 3
) -> List[Dict[str, Any]]:
    """
    Query a project's knowledge base for relevant context

    Args:
        project_id: Project identifier
        query: Query text
        n_results: Number of results to return

    Returns:
        List of relevant document chunks with metadata
    """
    try:
        collection = get_or_create_collection(project_id)

        results = collection.query(
            query_texts=[query],
            n_results=n_results
        )

        # Format results
        formatted_results = []
        if results and results['documents'] and len(results['documents']) > 0:
            documents = results['documents'][0]
            metadatas = results['metadatas'][0] if results['metadatas'] else [{}] * len(documents)
            distances = results['distances'][0] if results['distances'] else [0.0] * len(documents)

            for doc, meta, dist in zip(documents, metadatas, distances):
                formatted_results.append({
                    "content": doc,
                    "metadata": meta,
                    "similarity": 1.0 - dist  # Convert distance to similarity
                })

        return formatted_results

    except Exception as e:
        print(f"Error querying project knowledge: {e}")
        return []

def delete_project_knowledge(project_id: str) -> bool:
    """
    Delete all knowledge base data for a project

    Args:
        project_id: Project identifier

    Returns:
        True if successful
    """
    try:
        client = get_chroma_client()
        # Deliberately the same function `get_or_create_collection` builds the name with: a
        # second copy of the hash expression here is how a delete stops matching its create.
        name = collection_name(project_id)
        # Collection may already be absent — that counts as success.
        try:
            client.delete_collection(name=name)
        except ValueError:
            pass  # "Collection … does not exist" — nothing to delete
        return True
    except Exception as e:
        print(f"Error deleting project knowledge: {e}")
        return False

def get_project_document_count(project_id: str) -> int:
    """
    Get the number of documents in a project's knowledge base

    Args:
        project_id: Project identifier

    Returns:
        Number of documents
    """
    try:
        collection = get_or_create_collection(project_id)
        return collection.count()
    except BaseException:
        # BaseException catches pyo3_runtime.PanicException from ChromaDB's
        # Rust bindings (e.g. corrupted SQLite after version upgrade).
        return 0

# Text extraction utilities

def extract_text_from_file(file_path: Path) -> str:
    """
    Extract text from various file formats

    Args:
        file_path: Path to file

    Returns:
        Extracted text content
    """
    ext = file_path.suffix.lower()

    try:
        if ext == '.txt' or ext == '.md':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        elif ext == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    return text
            except ImportError:
                return f"[PDF file: {file_path.name} - PyPDF2 not installed for text extraction]"

        elif ext == '.docx':
            if not DOCX_AVAILABLE:
                return f"[DOCX file: {file_path.name} - python-docx not installed. pip install python-docx]"
            doc = _docx_lib.Document(str(file_path))
            parts: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)

        elif ext == '.xlsx':
            if not XLSX_AVAILABLE:
                return f"[XLSX file: {file_path.name} - openpyxl not installed. pip install openpyxl]"
            wb = _openpyxl_lib.load_workbook(str(file_path), read_only=True, data_only=True)
            parts: list[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"--- Sheet: {sheet_name} ---")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        parts.append(" | ".join(cells))
            wb.close()
            return "\n".join(parts)

        else:
            return f"[Unsupported file type: {ext}]"

    except Exception as e:
        return f"[Error extracting text from {file_path.name}: {str(e)}]"

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """
    Split text into overlapping chunks for better retrieval

    Args:
        text: Text to chunk
        chunk_size: Size of each chunk in characters
        overlap: Overlap between chunks

    Returns:
        List of text chunks
    """
    if not text or len(text) <= chunk_size:
        return [text] if text else []

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # Try to break at sentence boundary
        if end < len(text):
            # Look for period, question mark, or exclamation
            for i in range(end, start + chunk_size - 100, -1):
                if text[i] in '.!?\n':
                    end = i + 1
                    break

        chunks.append(text[start:end].strip())
        start = end - overlap

    return chunks

def process_and_add_file(project_id: str, file_path: Path) -> int:
    """
    Process a file and add it to the project's knowledge base

    Args:
        project_id: Project identifier
        file_path: Path to file

    Returns:
        Number of chunks added
    """
    # Extract text
    text = extract_text_from_file(file_path)

    if not text or text.startswith("["):
        return 0

    # Chunk text
    chunks = chunk_text(text)

    if not chunks:
        return 0

    # Generate metadata
    metadatas = [{
        "source": file_path.name,
        "chunk_index": i,
        "total_chunks": len(chunks)
    } for i in range(len(chunks))]

    # Generate IDs
    ids = [
        f"{hashlib.md5(file_path.name.encode()).hexdigest()[:8]}_{i}"
        for i in range(len(chunks))
    ]

    # Add to collection
    return add_documents_to_project(project_id, chunks, metadatas, ids)


def query_project_knowledge_filtered(
    project_id: str,
    query: str,
    n_results: int = 3,
    *,
    allowed_item_ids: Optional[List[str]] = None,
) -> List[Dict[str, Any]]:
    """
    Query project knowledge, optionally restricting to a set of project_items ids.

    This enables persona-scoped Chat-with-Docs:
    - persona attachments define allowed_item_ids
    - retrieval only searches those docs

    Backward compatible: if allowed_item_ids is None/empty, behaves like
    the standard query_project_knowledge.
    """
    try:
        collection = get_or_create_collection(project_id)

        where = None
        if allowed_item_ids:
            where = {"item_id": {"$in": allowed_item_ids}}

        results = collection.query(
            query_texts=[query],
            n_results=n_results,
            where=where,
        )

        formatted_results = []
        if results and results.get("documents") and len(results["documents"]) > 0:
            documents = results["documents"][0]
            metadatas = results["metadatas"][0] if results.get("metadatas") else [{}] * len(documents)
            distances = results["distances"][0] if results.get("distances") else [0.0] * len(documents)

            for doc, meta, dist in zip(documents, metadatas, distances):
                formatted_results.append({
                    "content": doc,
                    "metadata": meta,
                    "similarity": 1.0 - dist,
                })

        return formatted_results

    except Exception as e:
        print(f"Error querying project knowledge (filtered): {e}")
        return []


def process_and_add_file_with_item_id(
    project_id: str,
    file_path: Path,
    *,
    item_id: str,
    asset_id: str = "",
    original_name: str = "",
) -> int:
    """
    Process a file and add it to the project's knowledge base,
    tagging each chunk with item_id so we can do persona-scoped retrieval.
    """
    text = extract_text_from_file(file_path)
    if not text or text.startswith("["):
        return 0

    chunks = chunk_text(text)
    if not chunks:
        return 0

    source_name = original_name or file_path.name

    metadatas = [{
        "source": source_name,
        "chunk_index": i,
        "total_chunks": len(chunks),
        "item_id": item_id,
        "asset_id": asset_id,
    } for i in range(len(chunks))]

    ids = [
        f"{hashlib.md5((item_id + '_' + source_name).encode()).hexdigest()[:8]}_{i}"
        for i in range(len(chunks))
    ]

    return add_documents_to_project(project_id, chunks, metadatas, ids)
